"""S18: HTML/PDF/DOCX export from content_json."""
import html
import io
import json
from typing import Any

HORIZON_LABEL = {"short": "Kısa Vade", "medium": "Orta Vade", "long": "Uzun Vade"}


def _esc(val: Any) -> str:
    return html.escape(str(val) if val is not None else "")


def _commentary_html(section: dict) -> str:
    commentary = section.get("commentary", "")
    if not commentary:
        return ""
    body = _esc(commentary).replace("\n", "<br/>")
    return f'<p class="commentary">{body}</p>'


def render_section_html(section: dict) -> str:
    stype = section.get("type", "text")
    title = _esc(section.get("title", ""))

    if stype == "cover":
        subtitle = section.get("subtitle", "")
        sub_html = f'<p class="meta subtitle">{_esc(subtitle)}</p>' if subtitle else ""
        return f"""
        <div class="section cover">
          <h1>{_esc(section.get('title', ''))}</h1>
          {sub_html}
          <p class="meta">Müşteri: {_esc(section.get('client', ''))}</p>
          <p class="meta">Proje: {_esc(section.get('project', ''))}</p>
          <p class="meta">Tarih: {_esc(section.get('date', ''))}</p>
        </div>"""

    if stype == "divider":
        return '<hr class="divider"/>'

    if stype == "text":
        body = _esc(section.get("body", "")).replace("\n", "<br/>")
        return f'<div class="section"><h2>{title}</h2><p>{body}</p></div>'

    if stype == "kpi_grid":
        items = section.get("items", [])
        cells = "".join(
            f'<div class="kpi"><div class="kpi-val">{_esc(i.get("value"))}</div>'
            f'<div class="kpi-lbl">{_esc(i.get("label"))}</div></div>'
            for i in items
        )
        return f'<div class="section"><h2>{title}</h2><div class="kpi-grid">{cells}</div>{_commentary_html(section)}</div>'

    if stype == "table":
        cols = section.get("columns", [])
        rows = section.get("rows", [])
        thead = "".join(f"<th>{_esc(c)}</th>" for c in cols)
        tbody = "".join(
            "<tr>" + "".join(f"<td>{_esc(cell)}</td>" for cell in row) + "</tr>"
            for row in rows
        )
        return f'<div class="section"><h2>{title}</h2><table><thead><tr>{thead}</tr></thead><tbody>{tbody}</tbody></table></div>'

    if stype == "chart_radar":
        data = section.get("data", {})
        labels = data.get("labels", [])
        values = data.get("values", [])
        rows = "".join(
            f"<tr><td>{_esc(lbl)}</td><td>{_esc(val)}</td></tr>"
            for lbl, val in zip(labels, values)
        )
        return f'<div class="section"><h2>{title}</h2><table><thead><tr><th>Workstream</th><th>Skor</th></tr></thead><tbody>{rows}</tbody></table>{_commentary_html(section)}</div>'

    if stype == "chart_heatmap":
        data = section.get("data", {})
        cells = data.get("cells", [])
        rows = "".join(
            f"<tr><td>{_esc(c.get('capability_area'))}</td><td>{_esc(c.get('severity'))}</td>"
            f"<td>{_esc(c.get('risk_count'))}</td></tr>"
            for c in cells
        )
        return (
            f'<div class="section"><h2>{title}</h2>'
            f'<table><thead><tr><th>Alan</th><th>Severity</th><th>Sayı</th></tr></thead>'
            f"<tbody>{rows}</tbody></table>{_commentary_html(section)}</div>"
        )

    return f'<div class="section"><h2>{title}</h2><p>—</p></div>'


def content_to_html(report_title: str, content_json: str | None) -> str:
    data = json.loads(content_json) if content_json else {"sections": []}
    sections = data.get("sections", [])
    body = "".join(render_section_html(s) for s in sections)
    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/><title>{_esc(report_title)}</title>
<style>
  body {{ font-family: Arial, sans-serif; color: #1e293b; margin: 36px; line-height: 1.5; font-size: 13px; }}
  h1 {{ color: #1d4ed8; font-size: 22px; }}
  h2 {{ color: #334155; font-size: 16px; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px; margin-top: 24px; }}
  .meta {{ color: #64748b; font-size: 12px; }}
  table {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 6px 8px; text-align: left; }}
  th {{ background: #f1f5f9; font-size: 11px; }}
  .kpi-grid {{ display: flex; flex-wrap: wrap; gap: 12px; }}
  .kpi {{ border: 1px solid #e2e8f0; border-radius: 6px; padding: 10px 14px; min-width: 100px; }}
  .kpi-val {{ font-size: 20px; font-weight: bold; color: #1d4ed8; }}
  .kpi-lbl {{ font-size: 10px; color: #64748b; }}
  .divider {{ margin: 24px 0; border: none; border-top: 1px solid #e2e8f0; }}
</style></head><body>{body}
<p style="margin-top:32px;font-size:10px;color:#94a3b8;">AAKP — AI Assessment Knowledge Platform</p>
</body></html>"""


def export_pdf_bytes(html_doc: str) -> bytes:
    try:
        from xhtml2pdf import pisa
        buf = io.BytesIO()
        pisa.CreatePDF(html_doc, dest=buf, encoding="utf-8")
        return buf.getvalue()
    except Exception:
        # Fallback: return HTML as bytes if PDF engine unavailable
        return html_doc.encode("utf-8")


def export_docx_bytes(report_title: str, content_json: str | None) -> bytes:
    from docx import Document

    data = json.loads(content_json) if content_json else {"sections": []}
    doc = Document()
    doc.add_heading(report_title, 0)

    for section in data.get("sections", []):
        stype = section.get("type", "text")
        title = section.get("title", "")

        if stype == "cover":
            doc.add_heading(str(section.get("title", report_title)), level=1)
            if section.get("subtitle"):
                doc.add_paragraph(str(section.get("subtitle")))
            doc.add_paragraph(f"Müşteri: {section.get('client', '')}")
            doc.add_paragraph(f"Proje: {section.get('project', '')}")
            doc.add_paragraph(f"Tarih: {section.get('date', '')}")
            continue

        if stype == "divider":
            doc.add_paragraph("—" * 40)
            continue

        if title:
            doc.add_heading(title, level=2)

        if stype == "text":
            doc.add_paragraph(section.get("body", ""))
        elif stype == "kpi_grid":
            for item in section.get("items", []):
                doc.add_paragraph(f"{item.get('label')}: {item.get('value')}")
            if section.get("commentary"):
                doc.add_paragraph(str(section.get("commentary")))
        elif stype == "table":
            cols = section.get("columns", [])
            rows = section.get("rows", [])
            if cols:
                table = doc.add_table(rows=1, cols=len(cols))
                for i, c in enumerate(cols):
                    table.rows[0].cells[i].text = str(c)
                for row in rows:
                    cells = table.add_row().cells
                    for i, cell in enumerate(row):
                        if i < len(cells):
                            cells[i].text = str(cell)
        elif stype in ("chart_radar", "chart_heatmap"):
            if stype == "chart_radar":
                data = section.get("data", {})
                for lbl, val in zip(data.get("labels", []), data.get("values", [])):
                    doc.add_paragraph(f"{lbl}: {val}")
            else:
                for c in section.get("data", {}).get("cells", []):
                    doc.add_paragraph(
                        f"{c.get('capability_area')} / {c.get('severity')}: {c.get('risk_count')}"
                    )
            if section.get("commentary"):
                doc.add_paragraph(str(section.get("commentary")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
